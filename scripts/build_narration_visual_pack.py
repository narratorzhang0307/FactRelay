#!/usr/bin/env python3
"""Build a 36-shot Fact Atlas visual pack aligned to the 150-second narration."""

from __future__ import annotations

import csv
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_demo_video import (
    BLACK,
    CYAN,
    FONT_BOLD,
    FONT_CN,
    HEIGHT,
    LIME,
    PAPER,
    PINK,
    VIOLET,
    WIDTH,
    YELLOW,
    close_frame,
    font,
)
from build_storyboard_frames import (
    atlas_map_frame,
    council_frame,
    focus,
    gonka_proof_frame,
    gonka_router_frame,
    gonka_sequence_frame,
    handoff_frame,
    human_gate_frame,
    paths_frame,
    privacy_frame,
    product_frame,
    safety_frame,
    signal_receipt_frame,
    signals_frame,
    unplaced_frame,
)
from build_subtitle_master import SUBTITLES
from build_video_script import add_field, set_font, set_table_geometry, shade


ROOT = Path(__file__).resolve().parents[1]
DELIVERY = ROOT / "FactAtlas_黑客松交付包"
OUT = DELIVERY / "06_口播画面包_36张"
RAW = OUT / "00_真实界面原始截图"
FRAMES = OUT / "01_按时间码排列_36张"
OLD_SHOTS = ROOT / "FactRelay_黑客松交付包" / "03_截图"


# cue index, slug, source mode, source argument, purpose, phrase shown during this shot
FRAME_PLAN = [
    (0, "知识黑盒问题", "concept", "product", "概念开场", "个人知识库最危险的，不是缺少信息。"),
    (0, "未核验主张变成事实", "mobile", "M01_Relay手机_产品开场.png", "手机端产品展示", "真正的风险，是把未经验证的说法当成事实。"),
    (1, "网页与手机双端适配", "concept", "responsive", "双端响应式适配", "Fact Atlas 同时适配网页端和手机端。"),
    (1, "Relay_Atlas_Signals三Tab", "concept", "paths", "两条知识路径", "三个标签连接主动探索与每日发现。"),
    (2, "Relay桌面工作台", "desktop", "D02_Relay桌面_核验工作台.png", "真实桌面交互", "在 Relay，你可以从一条待核验主张开始。"),
    (2, "Relay手机三种输入", "mobile", "M02_Relay手机_核验工作台.png", "手机端输入交互", "文本、公开链接或截图，都能发起深度核验。"),
    (3, "八个全球主题Agent", "concept", "signals", "Signals 概念全景", "在 Signals，不是无限信息流，而是主题化的公开信号。"),
    (3, "Signals手机日期单卡", "mobile", "M08_Signals手机_日期单卡浏览.png", "Signals 手机端按日期单卡浏览", "选择主题与日期，一次只看一张可左右滑动的新闻卡。"),
    (4, "Gonka重要性排序", "concept", "signal_receipt", "Gonka 筛选概念", "第一道 Gonka 筛选，只判断一条信息是否值得关注。"),
    (4, "Signals手机筛选回执", "mobile", "M07_Signals手机_Gonka情报卡.png", "Signals 手机端 Gonka 情报卡", "它保留上游请求回执，但重要性绝不冒充真实度。"),
    (5, "Signals汇入Relay", "concept", "handoff", "用户控制的流转链", "候选信号不会自动进入知识星球。"),
    (5, "DeepVerify手机交互", "mobile", "M07_Signals手机_Gonka情报卡.png", "手机端 Deep verify 交互", "只有你选中它，才会进入 FactRelay 的证据检索与交叉验证。"),
    (6, "来源账本全景", "focus", ("06_真实结果_02_来源.png", (0.20, 0.04, 0.99, 0.94)), "真实来源账本", "每条来源都被放进一份编号账本。"),
    (6, "来源字段近景", "focus", ("06_真实结果_02_来源.png", (0.34, 0.22, 0.97, 0.86)), "发布者、日期、链接、立场与可信度", "发布者、日期、原始链接、立场和可信度都可回看。"),
    (7, "TruthScore真实结果", "focus", ("05_真实结果_01_结论.png", (0.21, 0.03, 0.99, 0.91)), "真实结论与评分", "最终结果展示可读的 Truth Score 和结论信心。"),
    (7, "TruthScore确定性公式", "concept", "score", "确定性评分公式", "分数来自可测试代码，而不是由模型随口给分。"),
    (8, "越界来源编号被拒绝", "focus", ("06_真实结果_02_来源.png", (0.31, 0.33, 0.99, 0.99)), "来源编号校验", "模型只能引用账本中真实存在的来源编号。"),
    (8, "分歧与信任边界", "concept", "safety", "信任边界概念图", "无效编号会被拒绝，两个模型的分歧也会原样保留。"),
    (9, "FactRelay主Agent六子Agent", "desktop", "agent_architecture_desktop.png", "主 Agent、六个子 Agent 与 Skills", "FactRelay 主 Agent 编排六个边界清晰的子 Agent。"),
    (9, "Agent编排手机卡片链", "mobile", "agent_architecture_mobile.png", "手机端可滑动 Agent 卡片链", "受理、证据、调查、质疑、裁决和回执逐步交接。"),
    (10, "Kimi桌面调查卡", "desktop", "D04_Council桌面_Kimi调查方.png", "真实 Kimi 卡片", "Kimi-K2.6 担任调查方。"),
    (10, "Kimi证据判断近景", "focus", ("07_真实结果_03_双模型审查.png", (0.32, 0.26, 0.67, 0.99)), "Kimi 调查方推理", "它先形成一份基于编号证据的调查判断。"),
    (11, "MiniMax手机审查界面", "mobile", "M03_Council手机_四方审理.png", "MiniMax 手机端审查", "MiniMax-M2.7 担任对抗式质疑方。"),
    (11, "MiniMax对抗审查近景", "focus", ("07_真实结果_03_双模型审查.png", (0.60, 0.26, 0.94, 0.99)), "MiniMax 质疑方推理", "它专门检查循环引用、时间错位与遗漏背景。"),
    (12, "GonkaRequestID真实回执", "focus", ("08_真实结果_04_Gonka回执.png", (0.20, 0.10, 0.99, 0.90)), "Proof 区块真实回执", "每次推理都保留 Gonka 上游 Request ID 和执行顺序。"),
    (12, "公开站真实运行证明", "concept", "gonka_proof", "生产环境证明", "这些非空回执证明分析来自哪次真实请求。"),
    (13, "GonkaRouter唯一推理边界", "concept", "gonka_router", "核心技术架构", "所有语义推理，都通过 GonkaRouter 这条可见边界。"),
    (13, "Gonka双模型执行序列", "concept", "gonka_sequence", "双模型时序架构", "但回执只证明调用来源，不会被冒充为“事实证明”。"),
    (14, "Nominatim候选与人工门", "concept", "human_gate", "人工落位概念", "最终由用户决定，哪些知识进入自己的星球。"),
    (14, "Atlas手机落位界面", "mobile", "M05_Atlas手机_用户确认落位.png", "真实手机落位交互", "地点候选也必须由用户亲自确认。"),
    (15, "Mapbox深色知识地球", "concept", "atlas_map", "Mapbox 知识地球概念", "Atlas 使用深色 Mapbox 地球，让地理退到背景。"),
    (15, "Mapbox手机实际展示", "mobile", "M04_Atlas手机_Mapbox知识地球.png", "真实手机 Mapbox", "亮色事实节点、Truth Score 与可解释连线，会成为主角。"),
    (16, "未落位轨道", "concept", "unplaced", "不伪造坐标概念", "有些知识本来就没有可靠的地理语义。"),
    (16, "NoFabricatedCoordinates界面", "mobile", "M04_Atlas手机_Mapbox知识地球.png", "手机端不伪造坐标提示", "它们会明确留在未落位轨道，系统绝不伪造坐标。"),
    (17, "浏览器本地私人知识谱系", "concept", "privacy", "隐私与知识谱系", "完整证据快照只保存在当前浏览器，形成可回看的私人知识谱系。"),
    (18, "FactAtlas品牌收尾", "close", None, "项目名与公开链接", "Fact Atlas：构建你的知识世界，让每条事实都保留回执。"),
]


ACCENTS = [VIOLET, LIME, CYAN, YELLOW, PINK]


def word_color(value: str) -> str:
    """Normalize CSS-style colors for WordprocessingML."""
    return value.removeprefix("#")


def timestamp_seconds(value: str) -> float:
    hours, minutes, remainder = value.split(":")
    seconds, millis = remainder.split(",")
    return int(hours) * 3600 + int(minutes) * 60 + int(seconds) + int(millis) / 1000


def format_timestamp(value: float) -> str:
    total_ms = round(value * 1000)
    hours, rest = divmod(total_ms, 3_600_000)
    minutes, rest = divmod(rest, 60_000)
    seconds, millis = divmod(rest, 1000)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{millis:03d}"


def fit_crop(source: Path, box: tuple[float, float, float, float] | None = None) -> Image.Image:
    image = Image.open(source).convert("RGB")
    if box:
        width, height = image.size
        image = image.crop((int(width * box[0]), int(height * box[1]), int(width * box[2]), int(height * box[3])))
    target_ratio = 16 / 9
    ratio = image.width / image.height
    if ratio > target_ratio:
        crop_width = int(image.height * target_ratio)
        left = max(0, (image.width - crop_width) // 2)
        image = image.crop((left, 0, left + crop_width, image.height))
    else:
        crop_height = int(image.width / target_ratio)
        top = max(0, min((image.height - crop_height) // 3, image.height - crop_height))
        image = image.crop((0, top, image.width, top + crop_height))
    return image.resize((WIDTH, HEIGHT), Image.Resampling.LANCZOS)


def desktop_frame(source: Path, label: str, crop: tuple[float, float, float, float] | None = None) -> Image.Image:
    shot = fit_crop(source, crop)
    canvas = Image.new("RGB", (WIDTH, HEIGHT), BLACK)
    draw = ImageDraw.Draw(canvas)
    shot = shot.resize((1760, 990), Image.Resampling.LANCZOS)
    canvas.paste(shot, (80, 60))
    draw.rounded_rectangle((80, 28, 660, 94), radius=24, fill=LIME, outline=BLACK, width=3)
    draw.text((112, 46), "LIVE PRODUCT · 真实产品界面", font=font(FONT_CN, 20), fill=BLACK)
    draw.rounded_rectangle((1320, 28, 1840, 94), radius=24, fill=PAPER, outline=BLACK, width=3)
    draw.text((1350, 47), label[:32], font=font(FONT_CN, 18), fill=BLACK)
    return canvas


def paste_phone(canvas: Image.Image, source: Path, center_x: int, top: int, screen_height: int = 790) -> tuple[int, int, int, int]:
    """Place one real mobile capture inside a restrained, centered phone shell."""
    screenshot = Image.open(source).convert("RGB")
    screen_width = round(screenshot.width * screen_height / screenshot.height)
    screenshot = screenshot.resize((screen_width, screen_height), Image.Resampling.LANCZOS)

    phone_width = screen_width + 34
    phone_height = screen_height + 54
    left = center_x - phone_width // 2
    right = left + phone_width
    bottom = top + phone_height
    draw = ImageDraw.Draw(canvas)
    draw.rounded_rectangle((left + 18, top + 18, right + 18, bottom + 18), radius=58, fill=VIOLET)
    draw.rounded_rectangle((left, top, right, bottom), radius=58, fill=PAPER, outline=BLACK, width=5)
    draw.rounded_rectangle((left + 9, top + 9, right - 9, bottom - 9), radius=50, fill=BLACK)

    screen_left = center_x - screen_width // 2
    screen_top = top + 27
    mask = Image.new("L", screenshot.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, screen_width, screen_height), radius=34, fill=255)
    canvas.paste(screenshot, (screen_left, screen_top), mask)
    speaker_width = 78
    draw.rounded_rectangle(
        (center_x - speaker_width // 2, top + 13, center_x + speaker_width // 2, top + 20),
        radius=4,
        fill="#4c4d48",
    )
    return left, top, right, bottom


def mobile_frame(source: Path, title: str, chinese: str) -> Image.Image:
    canvas = Image.new("RGB", (WIDTH, HEIGHT), BLACK)
    draw = ImageDraw.Draw(canvas)
    draw.rounded_rectangle((58, 46, 430, 112), radius=26, fill=LIME, outline=PAPER, width=2)
    draw.text((86, 66), "MOBILE VERSION · 手机版", font=font(FONT_CN, 22), fill=BLACK)

    title_font = font(FONT_CN, 25)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_width = title_box[2] - title_box[0]
    draw.text((WIDTH - 64 - title_width, 70), title, font=title_font, fill=PAPER)

    paste_phone(canvas, source, WIDTH // 2, 132, 780)
    cue_font = font(FONT_CN, 20)
    cue_box = draw.textbbox((0, 0), chinese, font=cue_font)
    cue_width = min(cue_box[2] - cue_box[0], WIDTH - 200)
    draw.rounded_rectangle((WIDTH // 2 - cue_width // 2 - 28, 995, WIDTH // 2 + cue_width // 2 + 28, 1048), radius=20, fill="#20211e")
    draw.text((WIDTH // 2 - cue_width // 2, 1008), chinese, font=cue_font, fill="#d7d9d3")
    return canvas


def responsive_frame() -> Image.Image:
    canvas = Image.new("RGB", (WIDTH, HEIGHT), PAPER)
    draw = ImageDraw.Draw(canvas)
    draw.text((86, 58), "RESPONSIVE PRODUCT · 双端适配", font=font(FONT_CN, 27), fill=VIOLET)
    draw.text((86, 122), "One product. Two native layouts.", font=font(FONT_BOLD, 61), fill=BLACK)
    draw.text((88, 211), "网页端与手机端，都是完整的 Fact Atlas。", font=font(FONT_CN, 36), fill="#555750")

    desktop = fit_crop(RAW / "D01_Relay桌面_产品开场.png").resize((1100, 619), Image.Resampling.LANCZOS)
    draw.rounded_rectangle((68, 328, 1208, 987), radius=38, fill=VIOLET)
    draw.rounded_rectangle((50, 310, 1190, 969), radius=38, fill=BLACK)
    canvas.paste(desktop, (70, 330))
    draw.rounded_rectangle((80, 336, 294, 394), radius=20, fill=CYAN, outline=BLACK, width=2)
    draw.text((108, 354), "DESKTOP · 网页版", font=font(FONT_CN, 18), fill=BLACK)

    paste_phone(canvas, RAW / "M01_Relay手机_产品开场.png", 1522, 218, 700)
    draw.rounded_rectangle((1344, 918, 1700, 978), radius=22, fill=LIME, outline=BLACK, width=2)
    draw.text((1380, 936), "MOBILE · 手机版", font=font(FONT_CN, 20), fill=BLACK)
    return canvas


def score_formula_frame() -> Image.Image:
    canvas = Image.new("RGB", (WIDTH, HEIGHT), PAPER)
    draw = ImageDraw.Draw(canvas)
    draw.text((86, 58), "DETERMINISTIC TRUTH SCORE · 确定性评分", font=font(FONT_CN, 28), fill=VIOLET)
    draw.text((86, 123), "The models do not grade themselves.", font=font(FONT_BOLD, 63), fill=BLACK)
    draw.text((88, 211), "模型不能给自己打分。", font=font(FONT_CN, 38), fill="#53554f")
    blocks = [
        (90, 370, 550, 760, LIME, "55%", "MODEL CONSENSUS", "双模型标准化共识"),
        (730, 370, 1190, 760, YELLOW, "45%", "SOURCE EVIDENCE", "来源可信度与立场加权"),
        (1370, 330, 1830, 800, PINK, "18", "TRUTH SCORE", "确定性代码输出\n0 — 100"),
    ]
    for x1, y1, x2, y2, accent, value, label, chinese in blocks:
        draw.rounded_rectangle((x1 + 14, y1 + 14, x2 + 14, y2 + 14), radius=38, fill=VIOLET)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=38, fill=accent, outline=BLACK, width=5)
        draw.text((x1 + 36, y1 + 46), value, font=font(FONT_BOLD, 82), fill=BLACK)
        draw.text((x1 + 36, y1 + 176), label, font=font(FONT_BOLD, 26), fill=BLACK)
        draw.multiline_text((x1 + 36, y1 + 242), chinese, font=font(FONT_CN, 24), fill="#3d3f3a", spacing=10)
    for x1, x2 in ((550, 730), (1190, 1370)):
        draw.line((x1 + 22, 565, x2 - 26, 565), fill=BLACK, width=8)
        draw.polygon([(x2 - 32, 547), (x2 - 32, 583), (x2 - 4, 565)], fill=BLACK)
    draw.rounded_rectangle((90, 875, 1830, 965), radius=28, fill=BLACK)
    draw.text((130, 904), "Truth Score = 50 + 50 × (55% model signal + 45% evidence signal)", font=font(FONT_BOLD, 25), fill=CYAN)
    return canvas


def concept_frame(name: str) -> Image.Image:
    builders = {
        "product": product_frame,
        "responsive": responsive_frame,
        "paths": paths_frame,
        "signals": signals_frame,
        "signal_receipt": signal_receipt_frame,
        "handoff": handoff_frame,
        "score": score_formula_frame,
        "safety": safety_frame,
        "council": council_frame,
        "gonka_proof": gonka_proof_frame,
        "gonka_router": gonka_router_frame,
        "gonka_sequence": gonka_sequence_frame,
        "human_gate": human_gate_frame,
        "atlas_map": atlas_map_frame,
        "unplaced": unplaced_frame,
        "privacy": privacy_frame,
    }
    return builders[name]()


def split_segments() -> list[tuple[str, str]]:
    counts = Counter(item[0] for item in FRAME_PLAN)
    seen: defaultdict[int, int] = defaultdict(int)
    segments: list[tuple[str, str]] = []
    for cue_index, *_ in FRAME_PLAN:
        start_text, end_text, _, _ = SUBTITLES[cue_index]
        start = timestamp_seconds(start_text)
        end = timestamp_seconds(end_text)
        part = seen[cue_index]
        total = counts[cue_index]
        segment_start = start + (end - start) * part / total
        segment_end = start + (end - start) * (part + 1) / total
        segments.append((format_timestamp(segment_start), format_timestamp(segment_end)))
        seen[cue_index] += 1
    return segments


def contact_sheet(paths: list[Path], rows: list[dict[str, str]]) -> Image.Image:
    sheet = Image.new("RGB", (3840, 2160), BLACK)
    draw = ImageDraw.Draw(sheet)
    cell_w, cell_h = 640, 360
    for index, (path, row) in enumerate(zip(paths, rows)):
        grid_row, col = divmod(index, 6)
        x, y = col * cell_w, grid_row * cell_h
        thumb = Image.open(path).convert("RGB").resize((640, 360), Image.Resampling.LANCZOS)
        sheet.paste(thumb, (x, y))
        draw.rectangle((x, y + 300, x + 640, y + 360), fill="#11120fe8")
        draw.rounded_rectangle((x + 16, y + 312, x + 72, y + 348), radius=14, fill=ACCENTS[index % len(ACCENTS)])
        draw.text((x + 29, y + 319), f"{index + 1:02d}", font=font(FONT_BOLD, 15), fill=BLACK)
        draw.text((x + 86, y + 316), f"{row['start'][3:8]} → {row['end'][3:8]}", font=font(FONT_BOLD, 16), fill=PAPER)
        draw.text((x + 272, y + 315), row["slug"][:18], font=font(FONT_CN, 15), fill=CYAN)
    return sheet


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "STHeiti"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("11120F")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal._element.rPr.rFonts.set(qn(key), "STHeiti")

    for name, size, color, before, after in (
        ("Title", 26, "11120F", 0, 6),
        ("Heading 1", 16, "7865FF", 18, 10),
        ("Heading 2", 13, "7865FF", 14, 7),
        ("Heading 3", 12, "3C49DA", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "STHeiti"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style._element.rPr.rFonts.set(qn(key), "STHeiti")

    header = section.header.paragraphs[0]
    set_font(header.add_run("FACT ATLAS  /  36-SHOT NARRATION VISUAL PACK  ·  36 镜头口播画面包"), 7.5, True, "5F645E", "STHeiti")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("AI³ GROWTH HACKATHON 2026   ·   "), 7.5, True, "5F645E", "STHeiti")
    add_field(footer, "PAGE")

    doc.core_properties.title = "Fact Atlas 36 镜头口播画面对应表"
    doc.core_properties.subject = "150 秒口播稿与 36 张交互、概念、架构及界面图的精确对应"
    doc.core_properties.author = "Fact Atlas"
    doc.core_properties.keywords = "Fact Atlas, Gonka, narration, storyboard, 36 shots, Mapbox"


def add_summary_table(doc: Document, rows: list[dict[str, str]]) -> None:
    doc.add_heading("36-shot timeline / 36 镜头时间轴", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for cell, label in zip(headers, ("#", "TIME / 时间", "VISUAL / 画面", "NARRATION SWITCH / 口播切点")):
        shade(cell, word_color(VIOLET))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 8, True, "FFFFFF", "STHeiti")
    set_repeat_header(table.rows[0])
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        if index % 2 == 0:
            for cell in cells:
                shade(cell, "F1F0EB")
        values = (f"{index + 1:02d}", f"{row['start'][3:8]}\n→ {row['end'][3:8]}", row["purpose"], row["phrase"])
        for cell, value in zip(cells, values):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if cell in (cells[0], cells[1]):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(value), 7.6 if cell != cells[3] else 7.2, cell == cells[0], "11120F", "STHeiti")
    set_table_geometry(table, [620, 1160, 2360, 5220])


def build_docx(path: Path, frame_paths: list[Path], rows: list[dict[str, str]]) -> None:
    doc = Document()
    configure_doc(doc)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    set_font(kicker.add_run("AI³ GROWTH HACKATHON 2026  ·  EDITING HANDOFF"), 8.5, True, word_color(VIOLET), "STHeiti")
    title = doc.add_paragraph(style="Title")
    set_font(title.add_run("Fact Atlas 36 镜头口播画面对应表"), 26, True, "11120F", "STHeiti")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_font(subtitle.add_run("150-second narration → 36 editing-ready visuals / 150 秒口播 → 36 张可直接剪辑画面"), 11, True, "5F645E", "STHeiti")

    metrics = doc.add_table(rows=1, cols=4)
    metrics.style = "Table Grid"
    for cell, (value, label, fill) in zip(metrics.rows[0].cells, [
        ("02:30", "总时长", "EDE9FF"),
        ("36", "主镜头", LIME),
        ("10", "手机主镜头", CYAN),
        ("19", "口播母句", YELLOW),
    ]):
        shade(cell, word_color(fill))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value + "\n"), 13, True, "11120F", "STHeiti")
        set_font(p.add_run(label), 8, True, "5F645E", "STHeiti")
    set_table_geometry(metrics, [2340, 2340, 2340, 2340])

    doc.add_paragraph()
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    shade(note.cell(0, 0), "11120F")
    p = note.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("使用方法 / EDITOR NOTE\n"), 8, True, word_color(LIME), "STHeiti")
    set_font(p.add_run("你按现有 19 句中文口播母稿录音即可。同一句内的两张图已均分时间，剪辑时按本表时间码切换；真人语速有轻微偏差时，以“口播切点”短句为对齐锚点。"), 10, False, "FFFFFF", "STHeiti")
    set_table_geometry(note, [9360])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_summary_table(doc, rows)

    for index in range(0, len(rows), 2):
        doc.add_section(WD_SECTION.NEW_PAGE)
        group = rows[index:index + 2]
        for offset, row in enumerate(group):
            item_index = index + offset
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(2 if offset == 0 else 12)
            heading.paragraph_format.space_after = Pt(6)
            heading.paragraph_format.keep_with_next = True
            set_font(heading.add_run(f"{item_index + 1:02d}  "), 13, True, word_color(VIOLET), "STHeiti")
            set_font(heading.add_run(f"{row['start']} → {row['end']}  ·  {row['slug']}"), 11.5, True, "11120F", "STHeiti")

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            left, right = table.rows[0].cells
            left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_img = left.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(0)
            run_img = p_img.add_run()
            run_img.add_picture(str(frame_paths[item_index]), width=Inches(2.85))
            shade(left, "11120F")
            shade(right, word_color(["EDE9FF", LIME, CYAN, YELLOW, PINK][item_index % 5]))

            p = right.paragraphs[0]
            p.paragraph_format.space_after = Pt(4)
            set_font(p.add_run("画面用途 / VISUAL\n"), 7.8, True, word_color(VIOLET), "STHeiti")
            set_font(p.add_run(row["purpose"]), 10.2, True, "11120F", "STHeiti")
            p2 = right.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            set_font(p2.add_run("口播切点 / SWITCH CUE\n"), 7.8, True, word_color(VIOLET), "STHeiti")
            set_font(p2.add_run(row["phrase"]), 9.5, False, "11120F", "STHeiti")
            p3 = right.add_paragraph()
            p3.paragraph_format.space_after = Pt(0)
            set_font(p3.add_run("完整母句 / FULL LINE\n"), 7.8, True, word_color(VIOLET), "STHeiti")
            set_font(p3.add_run(row["narration"]), 8.6, False, "3F413C", "STHeiti")
            set_table_geometry(table, [4200, 5160])
    doc.save(path)


def write_mapping(rows: list[dict[str, str]]) -> None:
    csv_path = OUT / "00_36镜头口播对应表.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=["index", "start", "end", "duration", "filename", "slug", "purpose", "phrase", "narration", "english"])
        writer.writeheader()
        writer.writerows(rows)

    lines = [
        "# Fact Atlas 36 镜头口播画面对应表",
        "",
        "36 张 1920×1080 主画面覆盖 150 秒口播；其中 10 张是手机竖屏居中的响应式产品镜头，另保留 20 张桌面/手机公网原始实拍。",
        "",
        "| # | 时间码 | 文件 | 画面用途 | 口播切点 |",
        "| ---: | --- | --- | --- | --- |",
    ]
    for row in rows:
        lines.append(f"| {int(row['index']):02d} | {row['start']} → {row['end']} | `{row['filename']}` | {row['purpose']} | {row['phrase']} |")
    lines.extend([
        "",
        "## 剪辑方法",
        "",
        "1. 将 `01_按时间码排列_36张` 按文件名排序导入剪辑软件。",
        "2. 直接使用本表开始/结束时间；总时长精确为 150 秒。",
        "3. 录音时仍按现有 19 句中文母稿完整朗读，不需要为 36 张图重写口播。",
        "4. 每句内的两张图已默认均分时间；真人声音交回后，再以“口播切点”短句微调切换位置。",
        "5. 相邻图优先用 6–12 帧交叉溶解或直切；不使用花哨转场。",
        "6. `00_真实界面原始截图` 保留了无统一包装的原始界面，需要更长停留或局部放大时可替换主画面。",
        "",
    ])
    (OUT / "00_36镜头口播对应表.md").write_text("\n".join(lines), encoding="utf-8")


def write_readme() -> None:
    content = """# Fact Atlas 36 镜头双端适配口播画面包

这个目录专门用于“先录口播，再合成视频”的工作流。

- `00_真实界面原始截图`：20 张公网真实产品截图，包含桌面、手机端、日期单卡与 Agent 编排。
- `01_按时间码排列_36张`：36 张 1920×1080 主画面，已按口播时间线排序；其中 10 张是手机竖屏居中、简洁手机外壳包装的响应式镜头。
- `00_36镜头总览_4K.png`：6×6 总览，用于快速检查整体节奏和色彩。
- `00_36镜头口播对应表.csv/.md`：机器可读和人工可读的时间轴。
- `FactAtlas_36镜头口播画面对应表.docx/.pdf`：供口播、剪辑与审阅的正式文档。
- `FactAtlas_新版口播稿_网页手机双端适配.docx/.pdf`：已明确加入网页端与手机端适配说明的 150 秒正式口播母稿。

口播第 2 句已明确说明“同时适配网页端和手机端”。录制完成后，只需提供一个完整音频文件，后期再根据 36 个时间段对齐画面。
"""
    (OUT / "README_使用说明.md").write_text(content, encoding="utf-8")


def main() -> None:
    if len(FRAME_PLAN) != 36:
        raise SystemExit("The narration visual pack must contain exactly 36 frames.")
    if not RAW.exists():
        raise SystemExit(f"Missing raw browser captures: {RAW}")

    FRAMES.mkdir(parents=True, exist_ok=True)
    for old in FRAMES.glob("*.png"):
        old.unlink()

    segments = split_segments()
    frame_paths: list[Path] = []
    rows: list[dict[str, str]] = []
    for index, (plan, segment) in enumerate(zip(FRAME_PLAN, segments), start=1):
        cue_index, slug, mode, source_arg, purpose, phrase = plan
        start, end = segment
        narration = SUBTITLES[cue_index][2]
        english = SUBTITLES[cue_index][3]
        filename = f"{index:02d}_{start.replace(':', '-').replace(',', '-')}_{slug}.png"
        destination = FRAMES / filename

        if mode == "concept":
            image = concept_frame(str(source_arg))
            image.save(destination, quality=95)
        elif mode == "desktop":
            image = desktop_frame(RAW / str(source_arg), purpose)
            image.save(destination, quality=95)
        elif mode == "desktop_crop":
            source_name, crop_box = source_arg
            image = desktop_frame(RAW / source_name, purpose, crop_box)
            image.save(destination, quality=95)
        elif mode == "mobile":
            image = mobile_frame(RAW / str(source_arg), purpose, phrase)
            image.save(destination, quality=95)
        elif mode == "focus":
            source_name, crop_box = source_arg
            image = focus(OLD_SHOTS / source_name, crop_box, ACCENTS[(index - 1) % len(ACCENTS)])
            image.save(destination, quality=95)
        elif mode == "close":
            close_frame(destination)
        else:
            raise ValueError(mode)

        duration = timestamp_seconds(end) - timestamp_seconds(start)
        row = {
            "index": str(index),
            "start": start,
            "end": end,
            "duration": f"{duration:.3f}",
            "filename": filename,
            "slug": slug,
            "purpose": purpose,
            "phrase": phrase,
            "narration": narration,
            "english": english,
        }
        frame_paths.append(destination)
        rows.append(row)

    contact_sheet(frame_paths, rows).save(OUT / "00_36镜头总览_4K.png", quality=95)
    write_mapping(rows)
    write_readme()
    build_docx(OUT / "FactAtlas_36镜头口播画面对应表.docx", frame_paths, rows)
    print(f"Generated {len(frame_paths)} aligned frames, a 4K contact sheet, mappings, and DOCX in {OUT}")


if __name__ == "__main__":
    main()
