#!/usr/bin/env python3
"""Build the bilingual FactRelay demo-video script as a polished Word document."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "11120F"
MUTED = "5F645E"
VIOLET = "7865FF"
VIOLET_SOFT = "EDE9FF"
LIME = "B8FF5C"
CYAN = "BCECF2"
PINK = "FFBED8"
YELLOW = "FFE38A"
PAPER = "F8F7F2"
WHITE = "FFFFFF"
BLUE = "3C49DA"

SCENES = [
    {
        "number": "01",
        "time": "00:00–00:12",
        "title": "Open with the relay deck / 卡组开场",
        "screen": "Show the full FactRelay hero. Move the cursor to the compact arrows and flip Sources → Challenge → Proof once. Let the colored card layers remain visible.",
        "narration": "AI 事实核查不应该要求我们再相信另一个黑盒。FactRelay 把每次核查拆成一叠可以翻看的证据区块。",
        "subtitles": [
            ("00:00.000 → 00:05.800", "AI fact checking should not replace one black box with another."),
            ("00:05.800 → 00:12.000", "FactRelay turns each investigation into a navigable evidence chain."),
        ],
        "color": VIOLET,
    },
    {
        "number": "02",
        "time": "00:12–00:27",
        "title": "Submit any public claim / 提交主张",
        "screen": "Show the Text, Link, and Image tabs. Click each tab once, return to Text, then choose the Great Wall starter.",
        "narration": "用户可以提交原始文本、公开文章链接或社交媒体截图。实时运行会先提取准确主张，再检索当前公开资料。",
        "subtitles": [
            ("00:12.000 → 00:19.500", "Submit text, a public link, or a social-media screenshot."),
            ("00:19.500 → 00:27.000", "A live run extracts the claim, then retrieves current public evidence."),
        ],
        "color": LIME,
    },
    {
        "number": "03",
        "time": "00:27–00:41",
        "title": "Start one live case / 启动实时核查",
        "screen": "Use the claim “The Great Wall of China is visible from the Moon with the naked eye.” Click Run verification and briefly show the live progress state.",
        "narration": "这里核查一条常见说法：人类能从月球上肉眼看到长城。FactRelay 先收集证据，再依次调用两位 Gonka 模型。",
        "subtitles": [
            ("00:27.000 → 00:33.800", "We test the familiar claim that the Great Wall is visible from the Moon."),
            ("00:33.800 → 00:41.000", "FactRelay collects evidence before calling two Gonka models in sequence."),
        ],
        "color": CYAN,
    },
    {
        "number": "04",
        "time": "00:41–00:58",
        "title": "Read the verdict, not just the number / 读取结论",
        "screen": "Land on Block 01. Hold on the Truth Score, verdict seal, confidence bar, and four score signals.",
        "narration": "结果是 9 分，事实不符。但这个数字不是模型随口生成的。它由可测试代码，根据模型共识和来源加权证据确定计算。",
        "subtitles": [
            ("00:41.000 → 00:48.500", "The verdict is Refuted, with a Truth Score of 9."),
            ("00:48.500 → 00:58.000", "The score is deterministic: model consensus plus source-weighted evidence."),
        ],
        "color": PINK,
    },
    {
        "number": "05",
        "time": "00:58–01:16",
        "title": "Inspect the evidence ledger / 检查证据账本",
        "screen": "Use the right arrow to open Block 02 Sources. Point to publisher, date, URL, stance, reliability, and source number. Open one public source only if it is fast and reliable.",
        "narration": "每条证据都保留发布者、日期、原始链接、立场和可信度。模型只能引用这份编号清单，伪造或越界的来源编号会被拒绝。",
        "subtitles": [
            ("00:58.000 → 01:07.000", "Every source keeps its publisher, date, URL, stance, and reliability."),
            ("01:07.000 → 01:16.000", "Models may cite only this numbered ledger; invalid source indexes are rejected."),
        ],
        "color": YELLOW,
    },
    {
        "number": "06",
        "time": "01:16–01:36",
        "title": "Compare adversarial models / 比较双模型",
        "screen": "Flip to Block 03 Review. Pause on the lime Kimi investigator card and violet MiniMax skeptic card, then show their reasoning bullets and confidence values.",
        "narration": "Kimi-K2.6 担任调查方，先形成证据判断；MiniMax-M2.7 担任质疑方，专门寻找循环引用、时间错位、因果跳跃和遗漏背景。分歧不会被隐藏。",
        "subtitles": [
            ("01:16.000 → 01:25.500", "Kimi-K2.6 investigates and forms the first evidence-based judgment."),
            ("01:25.500 → 01:36.000", "MiniMax-M2.7 challenges it for source laundering, chronology errors, and missing context."),
        ],
        "color": VIOLET,
    },
    {
        "number": "07",
        "time": "01:36–01:54",
        "title": "Keep the real receipts / 保留真实回执",
        "screen": "Flip to Block 04 Proof. Zoom only enough for the real non-null Gonka Request IDs to be readable. Follow the execution order once.",
        "narration": "每次推理都保留 Gonka 上游 Request ID 和执行顺序。回执证明哪次请求生成了分析，但不冒充事实本身的证明；事实仍由可检查的证据支持。",
        "subtitles": [
            ("01:36.000 → 01:44.500", "Every inference preserves its upstream Gonka Request ID and execution order."),
            ("01:44.500 → 01:54.000", "A receipt proves which call produced the analysis—not whether the claim is true."),
        ],
        "color": CYAN,
    },
    {
        "number": "08",
        "time": "01:54–02:10",
        "title": "Show the trust boundary / 展示可信边界",
        "screen": "Switch to the GitHub README architecture section or a prepared architecture graphic. Highlight Gonka-only inference, deterministic scoring, SSRF protection, and null preview receipts.",
        "narration": "所有语义推理只通过 GonkaRouter。检索不依赖其他 AI；评分是确定性代码；链接会经过 SSRF 防护；预览数据绝不伪造真实回执。",
        "subtitles": [
            ("01:54.000 → 02:02.000", "All semantic inference runs exclusively through GonkaRouter."),
            ("02:02.000 → 02:10.000", "Tested scoring, SSRF guards, and null preview receipts make the boundary explicit."),
        ],
        "color": LIME,
    },
    {
        "number": "09",
        "time": "02:10–02:24",
        "title": "Return to the complete chain / 回到完整链路",
        "screen": "Return to the product. Click the four block index buttons in order: Verdict → Sources → Review → Proof. End on the full card stack.",
        "narration": "从结论、来源、审查到回执，每一张卡都是可翻看、可追溯的证据区块。用户看到的不只是答案，而是一条能够复核的链路。",
        "subtitles": [
            ("02:10.000 → 02:17.000", "Verdict, Sources, Review, and Proof remain individually inspectable."),
            ("02:17.000 → 02:24.000", "The user receives a reviewable chain—not just an answer."),
        ],
        "color": YELLOW,
    },
    {
        "number": "10",
        "time": "02:24–02:30",
        "title": "Close / 收尾",
        "screen": "Return to the hero. Hold on the FactRelay name, tagline, public demo URL, and GitHub URL.",
        "narration": "FactRelay：质疑主张，保留回执。",
        "subtitles": [("02:24.000 → 02:30.000", "FactRelay. Question the claim. Keep the receipts.")],
        "color": PINK,
    },
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str = "Arial Unicode MS") -> None:
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().rFonts
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border(paragraph, color=INK, size=10, space=1, side="bottom") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_label_paragraph(doc: Document, label: str, text: str, fill: str = PAPER) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "\n")
    set_font(r, 7.5, True, VIOLET)
    r2 = p.add_run(text)
    set_font(r2, 10, False, INK)
    p.paragraph_format.line_spacing = 1.18


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal._element.rPr.rFonts.set(qn(key), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 10),
        ("Heading 1", 17, BLUE, 18, 9),
        ("Heading 2", 13, INK, 13, 6),
        ("Heading 3", 11, INK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style._element.rPr.rFonts.set(qn(key), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(4)
    r = hp.add_run("FACTRELAY  /  DEMO PRODUCTION BOOK  ·  演示视频制作手册")
    set_font(r, 7.5, True, MUTED)
    set_paragraph_border(hp, color="B7BAB5", size=6)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(4)
    r = fp.add_run("AI³ GROWTH HACKATHON 2026   ·   ")
    set_font(r, 7.5, True, MUTED)
    add_field(fp, "PAGE")

    doc.core_properties.title = "FactRelay 演示视频字幕与分镜"
    doc.core_properties.subject = "AI³ Growth Hackathon 2026 · Track 3"
    doc.core_properties.author = "FactRelay"
    doc.core_properties.keywords = "FactRelay, Gonka, demo video, subtitles, AI fact checking"


def build_doc(output: Path) -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("AI³ GROWTH HACKATHON 2026  ·  TRACK 3")
    set_font(r, 8.5, True, VIOLET)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("FactRelay")
    set_font(r, 31, True, INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(3)
    r = subtitle.add_run("2:30 Demo Video Script")
    set_font(r, 20, True, INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("演示视频字幕与分镜 · 中文旁白 + 英文字幕")
    set_font(r, 13, True, MUTED)

    metrics = doc.add_table(rows=1, cols=4)
    metrics.style = "Table Grid"
    set_table_geometry(metrics, [2340, 2340, 2340, 2340])
    metric_data = [
        ("02:30", "TARGET · 目标时长", VIOLET),
        ("1920 × 1080", "CANVAS · 画面", CYAN),
        ("30 FPS", "CAPTURE · 帧率", LIME),
        ("CN → EN", "VOICE / SUBS · 声音/字幕", YELLOW),
    ]
    for cell, (value, label, fill) in zip(metrics.rows[0].cells, metric_data):
        shade(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value + "\n")
        set_font(r, 15 if value != "1920 × 1080" else 12.5, True, INK)
        r2 = p.add_run(label)
        set_font(r2, 6.5, True, INK)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360])
    c = callout.cell(0, 0)
    shade(c, PAPER)
    p = c.paragraphs[0]
    r = p.add_run("RECORDING RULE · 录制原则\n")
    set_font(r, 8, True, VIOLET)
    r2 = p.add_run("Show one real Gonka run. Keep both upstream Request IDs readable. Never expose the API key or personal information.\n必须展示一次真实 Gonka 核查与两条可读回执；不得录入 API Key 或任何个人信息。")
    set_font(r2, 10.5, True, INK)

    doc.add_heading("Recording setup / 录制设置", level=1)
    bullets = [
        "Browser zoom at 100%; one clean tab; close password managers, email, and unrelated tabs. / 浏览器缩放 100%，仅保留干净标签页。",
        "Chinese narration at a calm 1.0× pace; English subtitles no more than two lines. / 中文旁白自然语速；英文字幕最多两行。",
        "Use direct cuts. The card transitions already provide the visual rhythm. / 使用直接切镜，卡片翻页即为主要视觉节奏。",
        "Required on screen: both model names, one public source, Truth Score, and real non-null Request IDs. / 必须展示双模型、公开来源、评分与真实回执。",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(item), 10, False, INK)

    doc.add_page_break()
    doc.add_heading("Scene-by-scene production script / 逐镜头制作脚本", level=1)
    intro = doc.add_paragraph("Read the Chinese narration exactly as written. English lines are the subtitle master and may be pasted into the editor without rewriting. / 中文旁白按原文录制；英文行即字幕母版，可直接进入剪辑软件。")
    intro.paragraph_format.space_after = Pt(10)

    for index, scene in enumerate(SCENES):
        if index in {2, 4, 6, 8}:
            doc.add_page_break()
        bar = doc.add_table(rows=1, cols=2)
        bar.style = "Table Grid"
        set_table_geometry(bar, [1680, 7680])
        shade(bar.cell(0, 0), scene["color"])
        shade(bar.cell(0, 1), PAPER)
        p = bar.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{scene["number"]}\n{scene["time"]}')
        set_font(r, 8.5, True, INK)
        p = bar.cell(0, 1).paragraphs[0]
        r = p.add_run(scene["title"])
        set_font(r, 12, True, INK)

        add_label_paragraph(doc, "Screen / 画面", scene["screen"], WHITE)
        add_label_paragraph(doc, "Chinese narration / 中文旁白", scene["narration"], VIOLET_SOFT)

        subs = doc.add_table(rows=len(scene["subtitles"]), cols=2)
        subs.style = "Table Grid"
        set_table_geometry(subs, [2130, 7230])
        for row_idx, (timecode, text) in enumerate(scene["subtitles"]):
            shade(subs.cell(row_idx, 0), INK)
            shade(subs.cell(row_idx, 1), WHITE)
            p1 = subs.cell(row_idx, 0).paragraphs[0]
            r1 = p1.add_run(timecode)
            set_font(r1, 7.3, True, WHITE)
            p2 = subs.cell(row_idx, 1).paragraphs[0]
            r2 = p2.add_run(text)
            set_font(r2, 9.2, False, INK)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    doc.add_page_break()
    doc.add_heading("English subtitle master / 英文字幕总表", level=1)
    p = doc.add_paragraph("Keep every cue on screen for the full interval. Use sentence case and a dark translucent background. / 每条字幕完整覆盖对应区间，使用 sentence case 与深色半透明底。")
    p.paragraph_format.space_after = Pt(8)
    all_subs = [subtitle for scene in SCENES for subtitle in scene["subtitles"]]

    def add_subtitle_table(cues: list[tuple[str, str]]) -> None:
        table = doc.add_table(rows=1 + len(cues), cols=2)
        table.style = "Table Grid"
        set_table_geometry(table, [2130, 7230])
        for i, text in enumerate(("TIMECODE · 时间码", "ENGLISH SUBTITLE · 英文字幕")):
            shade(table.cell(0, i), INK)
            r = table.cell(0, i).paragraphs[0].add_run(text)
            set_font(r, 7.5, True, WHITE)
        for row_idx, (timecode, text) in enumerate(cues, start=1):
            shade(table.cell(row_idx, 0), VIOLET_SOFT if row_idx % 2 else PAPER)
            shade(table.cell(row_idx, 1), WHITE)
            set_font(table.cell(row_idx, 0).paragraphs[0].add_run(timecode), 7.5, True, INK)
            set_font(table.cell(row_idx, 1).paragraphs[0].add_run(text), 8.6, False, INK)

    add_subtitle_table(all_subs[:10])
    doc.add_page_break()
    doc.add_heading("English subtitle master — continued / 英文字幕续表", level=1)
    p = doc.add_paragraph("Cues 11–19 complete the same uninterrupted 2:30 timeline. / 第 11–19 条延续同一条完整的 2 分 30 秒时间线。")
    p.paragraph_format.space_after = Pt(8)
    add_subtitle_table(all_subs[10:])

    doc.add_page_break()
    doc.add_heading("Pronunciation and on-screen terms / 发音与术语", level=1)
    terms = [
        ("FactRelay", "“Fact Relay”; two English words spoken naturally / 两个英文词自然连读"),
        ("Gonka", "Keep the product name visible while it is spoken / 发音时画面保留 Gonka 名称"),
        ("Kimi-K2.6", "“Kimi K two point six”"),
        ("MiniMax-M2.7", "“MiniMax M two point seven”"),
        ("Truth Score", "Chinese narration may say 真实度评分; keep the English UI term visible"),
        ("Request ID", "Chinese narration may say 请求回执; keep the real ID readable"),
    ]
    term_table = doc.add_table(rows=len(terms), cols=2)
    term_table.style = "Table Grid"
    set_table_geometry(term_table, [2200, 7160])
    for idx, (term, note) in enumerate(terms):
        shade(term_table.cell(idx, 0), [LIME, CYAN, YELLOW, VIOLET_SOFT, PINK, PAPER][idx])
        set_font(term_table.cell(idx, 0).paragraphs[0].add_run(term), 9.5, True, INK)
        set_font(term_table.cell(idx, 1).paragraphs[0].add_run(note), 9.2, False, INK)

    doc.add_heading("Final recording and privacy check / 最终录制与隐私检查", level=1)
    checklist = [
        "Use a real live result; do not record only the preview fixture. / 使用真实运行结果，不只录预览数据。",
        "Both Gonka Request IDs are non-null and readable. / 两条 Gonka 回执非空且可读。",
        "No API key, .env.local, email, phone, password manager, or unrelated tab is visible. / 不出现密钥与个人信息。",
        "Pause 1–2 seconds after each card transition. / 每次翻卡后停留 1–2 秒。",
        "Watch the exported MP4 once with sound and once muted. / 导出后分别有声、静音完整检查。",
        "Test the public/unlisted link from a signed-out window. / 在未登录窗口测试公开视频链接。",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item), 10, False, INK)

    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    set_table_geometry(note, [9360])
    shade(note.cell(0, 0), YELLOW)
    p = note.cell(0, 0).paragraphs[0]
    r = p.add_run("FINAL EDITOR NOTE · 最终剪辑提示\n")
    set_font(r, 8, True, INK)
    r2 = p.add_run("Use cuts rather than decorative transitions. Add one quiet click at the first card flip only. Keep music at least 18 dB below narration—or omit it. Never speed up the score or Request ID shots beyond readability.\n以直接切镜为主，仅在第一次翻卡时加入轻微点击音；背景音乐至少比旁白低 18 dB，或不使用。评分与回执镜头必须保持可读。")
    set_font(r2, 9.7, False, INK)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_doc(args.output)
